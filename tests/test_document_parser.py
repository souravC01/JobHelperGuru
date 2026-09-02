import io
import docx
import pypdf
from backend.services.document_parser import extract_text_from_file

def test_extract_from_text():
    content = b"Senior Software Engineer\nPython, Docker, SQL"
    text = extract_text_from_file(content, "resume.txt")
    assert "Senior Software Engineer" in text
    assert "Python" in text

def test_extract_from_docx():
    doc = docx.Document()
    doc.add_heading("John Doe - Resume", level=1)
    doc.add_paragraph("Full Stack Developer with 5 years experience in React and Node.js.")
    doc.add_paragraph("Skills: TypeScript, PostgreSQL, Docker.")
    
    bio = io.BytesIO()
    doc.save(bio)
    docx_bytes = bio.getvalue()
    
    text = extract_text_from_file(docx_bytes, "my_resume.docx")
    assert "John Doe" in text
    assert "React" in text
    assert "PostgreSQL" in text

def test_extract_from_pdf():
    # Generate minimal valid PDF in-memory with pypdf
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    bio = io.BytesIO()
    writer.write(bio)
    pdf_bytes = bio.getvalue()
    
    text = extract_text_from_file(pdf_bytes, "sample.pdf")
    # Even on blank page, should return string without crashing
    assert isinstance(text, str)
