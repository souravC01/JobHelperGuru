import io
from pathlib import Path
from typing import Optional

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extracts plain text from PDF, Word (.docx), Markdown, or plain text files.
    """
    ext = Path(filename).suffix.lower()

    # 1. PDF
    if ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            extracted_pages = []
            for page in reader.pages:
                text = page.extract_text() or ""
                if text.strip():
                    extracted_pages.append(text.strip())
            return "\n\n".join(extracted_pages)
        except Exception as e:
            raise ValueError(f"Failed to read PDF file: {str(e)}")

    # 2. Word (.docx)
    if ext in [".docx", ".doc"]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())
            # Also extract text from any tables in the resume
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        lines.append(row_text)
            return "\n\n".join(lines)
        except Exception as e:
            # If docx fails (e.g. legacy binary .doc format), try reading strings as fallback
            try:
                decoded = file_bytes.decode("latin-1", errors="ignore")
                printable = "".join(c if 31 < ord(c) < 127 or c in "\n\r\t" else " " for c in decoded)
                cleaned = "\n".join(line.strip() for line in printable.splitlines() if len(line.strip()) > 3)
                if len(cleaned) > 50:
                    return cleaned
            except Exception:
                pass
            raise ValueError(f"Failed to read Word document: {str(e)}. Please save as .docx or .pdf.")

    # 3. Plain Text / Markdown / Other
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception as e:
            return file_bytes.decode("utf-8", errors="ignore")
